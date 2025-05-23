"""
Email Agent Implementation.

This module implements the core functionality for generating personalized
emails based on reasoning agent outputs.
"""

import os
import re
import sys
import logging
import json
import openai
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
import docx
from docx import Document
from docx.shared import Pt, Inches
import base64
from io import BytesIO
import platform
import msglib  # You'll need to install this: pip install msglib

from backend.email_agent.config import (
    MODEL_ID, TEMPERATURE, MAX_OUTPUT_TOKENS, API_KEY,
    REASONING_DIR, OUTPUT_DIR, BUYER_PREFIX, EMAIL_TEMPLATE_PATH,
    MAX_EMAILS_TO_GENERATE
)
from backend.email_agent.prompts import EMAIL_PROMPT

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('email_agent.log')
    ]
)
logger = logging.getLogger(__name__)

# Do not import Windows-specific modules on non-Windows platforms
if platform.system() == "Windows":
    import win32com.client
    import pythoncom
    import comtypes.client
else:
    # Flag to indicate Windows modules are unavailable
    HAS_WIN32COM = False

class EmailAgent:
    """
    Agent for generating personalized emails based on reasoning outputs.
    """
    
    def __init__(self, template_path=None, output_dir=None):
        """Initialize the email agent."""
        if not API_KEY:
            raise ValueError("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
        
        # Configure OpenAI client
        self.client = openai.OpenAI(api_key=API_KEY)
        
        # Use custom template path if provided
        self.template_path = template_path if template_path else EMAIL_TEMPLATE_PATH
        logger.info(f"Using template path: {self.template_path}")
        
        # Use custom output directory if provided
        self.output_dir = output_dir if output_dir else OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create No_Contact_Info subdirectory
        self.no_contact_dir = self.output_dir / "No_Contact_Info"
        self.no_contact_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created No_Contact_Info directory at: {self.no_contact_dir}")
        
        logger.info(f"Initialized EmailAgent with model: {MODEL_ID}")
    
    def list_strong_buyer_files(self) -> List[Dict[str, Any]]:
        """
        List all buyer files (STRONG, MEDIUM, and NOT) in the reasoning output directory.
        
        Returns:
            List of dictionaries with file information
        """
        files = []
        
        try:
            logger.info(f"Looking for files in: {REASONING_DIR}")
            logger.info(f"Directory exists: {REASONING_DIR.exists()}")
            logger.info(f"Directory is a directory: {REASONING_DIR.is_dir()}")
            
            # List all files in the directory for debugging
            all_files = list(REASONING_DIR.glob("*"))
            logger.info(f"All files in directory ({len(all_files)}): {[f.name for f in all_files]}")
            
            # Get all buyer files in the reasoning output directory (STRONG, MEDIUM, and NOT)
            # Create patterns for all buyer types
            patterns = ["STRONG*.txt", "MEDIUM*.txt", "NOT*.txt"] 
            
            all_matching_files = []
            for pattern in patterns:
                logger.info(f"Looking for pattern: {pattern}")
                matching_files = list(REASONING_DIR.glob(pattern))
                logger.info(f"Found {len(matching_files)} files matching {pattern}")
                all_matching_files.extend(matching_files)
            
            logger.info(f"Found {len(all_matching_files)} total buyer files")
            
            # Process all found files
            for file_path in all_matching_files:
                # Extract company name from filename
                filename = file_path.name
                
                # Extract buyer type (STRONG, MEDIUM, NOT)
                buyer_type = filename.split("_")[0]
                
                # Extract company name
                company_name_match = re.search(r'(?:STRONG|MEDIUM|NOT)_(.+?)_reasoning_', filename, re.IGNORECASE)
                
                if company_name_match:
                    company_name = company_name_match.group(1)
                    
                    files.append({
                        "path": str(file_path),  # Convert Path to string
                        "company_name": company_name,
                        "filename": filename,
                        "buyer_type": buyer_type,  # Add buyer type for reference
                        "processed": False,
                        "success": None,
                        "output_file": None,
                        "error": None
                    })
            
            logger.info(f"Found {len(files)} buyer files to process")
            
            # Limit the number of files if needed
            if len(files) > MAX_EMAILS_TO_GENERATE:
                logger.warning(f"Limiting to {MAX_EMAILS_TO_GENERATE} files")
                files = files[:MAX_EMAILS_TO_GENERATE]
                
            return files
            
        except Exception as e:
            logger.error(f"Error listing buyer files: {str(e)}")
            return []
    
    def extract_buyer_info(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract relevant buyer information from a reasoning file.
        
        Args:
            file_path: Path to the reasoning file
            
        Returns:
            Dictionary with extracted information
        """
        try:
            with open(file_path, 'r') as f:
                content = f.read()
            
            # Extract company name
            company_name_match = re.search(r'Company:\s*(.+?)\n', content)
            company_name = company_name_match.group(1) if company_name_match else "Unknown Company"
            
            # Extract URL - updated to be more reliable
            url_match = re.search(r'URL:\s*([^\s]+)', content)
            url = url_match.group(1).strip() if url_match else f"{company_name}.com"
            
            # If URL is somehow wrong, try to fix it
            if not url or url == "4.5M" or not re.match(r'^[a-zA-Z0-9][-a-zA-Z0-9.]*\.[a-zA-Z]{2,}$', url):
                url = f"{company_name}.com"
            
            # Extract contact information using multiple patterns
            contacts = []
            
            # Try to find contacts in the CONTACT INFORMATION section
            contact_section = re.search(r'CONTACT INFORMATION\n=+\n(.*?)\n=+', content, re.DOTALL)
            if contact_section and "No contact information found" not in contact_section.group(1):
                contact_text = contact_section.group(1)
                contact_lines = [line.strip() for line in contact_text.split('\n') if line.strip()]
                contacts.extend(contact_lines)
            
            # Also look for Key Team Members section with emails
            team_section = re.search(r'Key Team Members[:\s]*\n(.*?)(?:\n\n|\n=+)', content, re.DOTALL)
            if team_section:
                team_text = team_section.group(1).strip()
                team_lines = [line.strip() for line in team_text.split('\n') if line.strip() and '|' in line]
                contacts.extend(team_lines)
            
            # Look for emails in the entire document if we still don't have any
            if not contacts:
                email_matches = re.findall(r'([a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+)', content)
                if email_matches:
                    unique_emails = list(set(email_matches))
                    name_email_pattern = r'([A-Za-z\s.]+)(?:\([^)]+\))?\s*\|?\s*Email:\s*([a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+)'
                    name_email_matches = re.findall(name_email_pattern, content)
                    
                    if name_email_matches:
                        contacts = [f"{name.strip()} | Email: {email.strip()}" for name, email in name_email_matches]
                    else:
                        # Just list the emails if we can't associate names
                        contacts = [f"Contact: {email}" for email in unique_emails]
            
            # Extract reasoning from final assessment
            final_assessment = ""
            final_section = re.search(r'FINAL ASSESSMENT\n-+\n(.*?)\n(?:=+|CONTACT)', content, re.DOTALL)
            if final_section:
                final_assessment = final_section.group(1).strip()
            
            # Extract any other relevant sections for the email
            analysis_process = ""
            analysis_section = re.search(r'ANALYSIS PROCESS\n-+\n(.*?)\nFINAL ASSESSMENT', content, re.DOTALL)
            if analysis_section:
                analysis_process = analysis_section.group(1).strip()
            
            return {
                "company_name": company_name,
                "url": url,
                "contacts": contacts,
                "final_assessment": final_assessment,
                "analysis_process": analysis_process,
                "full_content": content
            }
            
        except Exception as e:
            logger.error(f"Error extracting buyer info from {file_path}: {str(e)}")
            return {
                "company_name": "Error",
                "url": "Error",
                "contacts": [],
                "final_assessment": "",
                "analysis_process": "",
                "full_content": f"Error: {str(e)}"
            }
    
    def encode_image(self, image_path: Path) -> str:
        """
        Encode an image to base64 for the Gemini API.
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Base64 encoded image data
        """
        if not image_path.exists():
            logger.error(f"Image file not found: {image_path}")
            return ""
        
        try:
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            logger.error(f"Error encoding image: {str(e)}")
            return ""
    
    def generate_email(self, buyer_info: Dict[str, Any]) -> str:
        """
        Generate a personalized email using the OpenAI model.
        
        Args:
            buyer_info: Dictionary with buyer information
            
        Returns:
            Generated email content
        """
        try:
            # Prepare content for the prompt
            prompt_text = EMAIL_PROMPT
            
            # Add template content
            if self.template_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.gif']:
                # Can't directly include images with text-only OpenAI models
                prompt_text += "\n\nNote: Please generate the email based on the template description provided in the prompt."
            else:
                # Handle text template
                try:
                    with open(self.template_path, 'r') as f:
                        template_text = f.read()
                    prompt_text += f"\n\nEmail Template:\n\n{template_text}"
                except Exception as e:
                    logger.error(f"Error reading template file: {e}")
                    raise ValueError(f"Failed to read template file at {self.template_path}")
            
            # Add buyer information
            prompt_text += f"\n\nBuyer profile information:\n{buyer_info['full_content']}"
            
            # Generate response using OpenAI
            logger.info(f"Generating email for {buyer_info['company_name']}")
            response = self.client.chat.completions.create(
                model=MODEL_ID,
                messages=[
                    {"role": "system", "content": "You are an expert email writer for investment banking"},
                    {"role": "user", "content": prompt_text}
                ],
                temperature=TEMPERATURE,
                max_tokens=MAX_OUTPUT_TOKENS
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ValueError("Empty response from OpenAI API")
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error generating email: {str(e)}")
            return f"Error generating email: {str(e)}"
    
    def save_email_as_docx(self, email_content: str, buyer_info: Dict[str, Any], has_contacts: bool, buyer_type: str = "UNKNOWN") -> Optional[Path]:
        """
        Save generated email as a Word document with proper formatting.
        
        Args:
            email_content: The generated email content
            buyer_info: Buyer information dictionary
            has_contacts: Whether contact information was found
            buyer_type: Type of buyer (STRONG, MEDIUM, NOT)
            
        Returns:
            Path to the saved document or None if failed
        """
        try:
            # Create a new document
            doc = Document()
            
            # Add metadata section at top
            doc.add_heading("META INFORMATION", level=1)
            
            # Add URL with proper formatting
            url_para = doc.add_paragraph()
            url_para.add_run("URL: ").bold = True
            
            # Ensure URL is not "4.5M" and is a valid domain
            url = buyer_info['url']
            if not url or url == "4.5M" or not re.match(r'^[a-zA-Z0-9][-a-zA-Z0-9.]*\.[a-zA-Z]{2,}$', url):
                url = f"{buyer_info['company_name']}.com"
            
            url_para.add_run(url)
            
            # Add team members with proper formatting and bullet points
            team_para = doc.add_paragraph()
            team_para.add_run("Team Members:").bold = True
            
            # Process contact information
            if buyer_info['contacts']:
                for contact in buyer_info['contacts']:
                    contact_para = doc.add_paragraph(style='ListBullet')
                    contact_para.add_run(contact)
            else:
                # Extract team information from the generated email
                team_section = re.search(r'\*\*Team Members:\*\*(.*?)(?:\*\*Subject:|$)', email_content, re.DOTALL)
                if team_section:
                    team_text = team_section.group(1).strip()
                    
                    # Process each team member line
                    for line in team_text.split('\n'):
                        if line.strip() and (line.strip().startswith('*') or line.strip().startswith('•')):
                            contact_para = doc.add_paragraph(style='ListBullet')
                            contact_para.add_run(line.strip().replace('* ', '').replace('• ', ''))
            
            # Add horizontal line
            doc.add_paragraph("─" * 50)
            
            # Extract and format the actual email content
            email_body = email_content
            
            # Remove the metadata section if it exists in the email content
            url_pattern = r'\*\*URL:\*\*.*?\n'
            team_pattern = r'\*\*Team Members:\*\*.*?(?=\*\*Subject:|$)'
            
            email_body = re.sub(url_pattern, '', email_body, flags=re.DOTALL)
            email_body = re.sub(team_pattern, '', email_body, flags=re.DOTALL)
            
            # Extract subject line
            subject_match = re.search(r'\*\*Subject:(.*?)\*\*', email_body)
            if subject_match:
                subject = subject_match.group(1).strip()
                email_body = re.sub(r'\*\*Subject:.*?\*\*', '', email_body, flags=re.DOTALL)
            else:
                subject = "Project Elevate: Premier Parking Lift Distributor Buyout Opportunity"
            
            # Add the subject as a heading
            doc.add_heading(subject, level=1)
            
            # Split into paragraphs
            paragraphs = re.split(r'\n\s*\n', email_body.strip())
            
            for paragraph in paragraphs:
                # Handle bullet point sections differently
                if '•' in paragraph or '*' in paragraph:
                    # Split by newlines to process each bullet point 
                    lines = paragraph.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                            
                        if line.startswith('•') or line.startswith('*'):
                            # This is a bullet point
                            bullet_para = doc.add_paragraph(style='ListBullet')
                            
                            # Extract the bullet text without the bullet symbol
                            bullet_text = re.sub(r'^[•*]\s*', '', line)
                            
                            # Look for bold text patterns and apply formatting
                            if '**' in bullet_text:
                                # Process bold text within the bullet point
                                parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        # This is bold text
                                        bold_text = part.strip('*')
                                        run = bullet_para.add_run(bold_text)
                                        run.bold = True
                                    else:
                                        bullet_para.add_run(part)
                            else:
                                # No formatting needed
                                bullet_para.add_run(bullet_text)
                        else:
                            # Regular text within a bullet point section
                            doc.add_paragraph(line)
                else:
                    # Regular paragraph
                    p = doc.add_paragraph()
                    
                    # Look for bold text with ** markers
                    if '**' in paragraph:
                        parts = re.split(r'(\*\*.*?\*\*)', paragraph)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                # This is bold text
                                bold_text = part.strip('*')
                                run = p.add_run(bold_text)
                                run.bold = True
                            else:
                                p.add_run(part)
                    else:
                        # No formatting needed
                        p.add_run(paragraph)
            
            # Save document with modified filename and directory
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            company_name = buyer_info['company_name'].replace('/', '_').replace('\\', '_')
            
            # Determine output directory based on contact info presence
            output_dir = self.no_contact_dir if not has_contacts else self.output_dir
            
            # Use buyer type in filename
            output_file = output_dir / f"{buyer_type}_Email_{company_name}_{timestamp}.docx"
            
            doc.save(output_file)
            logger.info(f"Saved email to {output_file}")
            
            return output_file
            
        except Exception as e:
            logger.error(f"Error saving email as Word document: {str(e)}")
            return None
    
    def save_email_as_outlook_msg(self, email_content: str, buyer_info: Dict[str, Any], has_contacts: bool, buyer_type: str = "UNKNOWN") -> Optional[Path]:
        """
        Save generated email as an Outlook .msg file on any platform.
        
        Args:
            email_content: The generated email content
            buyer_info: Buyer information dictionary
            has_contacts: Whether contact information was found
            buyer_type: Type of buyer (STRONG, MEDIUM, NOT)
            
        Returns:
            Path to the saved message file or None if failed
        """
        try:
            logger.info(f"Creating .eml file for {buyer_info['company_name']}")
            
            # Extract subject line
            subject = "Project Elevate: Premier Parking Lift Distributor Buyout Opportunity"  # Default
            subject_patterns = [
                r'\*\*Subject:(.*?)\*\*',  # Markdown format
                r'Subject:(.*?)(?:\n|$)',  # Plain text format
                r'Subject Line:(.*?)(?:\n|$)'  # Alternative format
            ]
            
            for pattern in subject_patterns:
                subject_match = re.search(pattern, email_content, re.IGNORECASE)
                if subject_match:
                    potential_subject = subject_match.group(1).strip()
                    if potential_subject:  # Ensure we have actual content
                        subject = potential_subject
                        break
            
            logger.info(f"Setting email subject: '{subject}'")
            
            # Extract and format recipients
            recipients = []
            recipient_emails = set()  # Using set to avoid duplicates
            
            # Extract from buyer_info contacts with improved regex
            if buyer_info['contacts']:
                logger.info(f"Processing {len(buyer_info['contacts'])} contact entries")
                
                for contact in buyer_info['contacts']:
                    # Various email patterns
                    email_matches = []
                    
                    # Standard email pattern
                    std_emails = re.findall(r'([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})', contact)
                    email_matches.extend(std_emails)
                    
                    # "Email:" label pattern
                    labeled_emails = re.findall(r'Email:\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})', contact, re.IGNORECASE)
                    email_matches.extend(labeled_emails)
                    
                    for email in email_matches:
                        email = email.strip()
                        if email and email not in recipient_emails and self._is_valid_email(email):
                            recipient_emails.add(email)
                            recipients.append(email)
                            logger.debug(f"Found recipient: {email}")
            
            # If no recipients found, search email content
            if not recipients:
                logger.info("No recipients found in contacts, searching email content")
                team_section = re.search(r'\*\*Team Members:\*\*(.*?)(?:\*\*Subject:|$)', email_content, re.DOTALL)
                
                if team_section:
                    team_text = team_section.group(1).strip()
                    
                    # Look for emails in team section
                    team_emails = re.findall(r'([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})', team_text)
                    
                    for email in team_emails:
                        email = email.strip()
                        if email and email not in recipient_emails and self._is_valid_email(email):
                            recipient_emails.add(email)
                            recipients.append(email)
                            logger.debug(f"Found recipient in team section: {email}")
            
            recipients_str = "; ".join(recipients) if recipients else ""
            logger.info(f"Recipients: {recipients_str or 'None'}")
            
            # Format email body - clean up markdown
            body = email_content
            
            # Remove metadata sections
            metadata_patterns = [
                (r'\*\*URL:\*\*.*?(?=\n\n|\n\*\*|\Z)', ''),  # URL section
                (r'\*\*Team Members:\*\*.*?(?=\n\n\*\*|\*\*Subject:|\Z)', ''),  # Team Members section
                (r'\*\*Subject:.*?(?=\n\n|\Z)', '')  # Subject line
            ]
            
            for pattern, replacement in metadata_patterns:
                body = re.sub(pattern, replacement, body, flags=re.DOTALL)
            
            # Remove markdown formatting
            body = re.sub(r'\*\*(.*?)\*\*', r'\1', body)  # Remove bold markers
            body = re.sub(r'^\s*[•*]\s*', '- ', body, flags=re.MULTILINE)  # Format bullet points
            body = re.sub(r'\n{3,}', '\n\n', body)  # Normalize line breaks
            body = re.sub(r' +$', '', body, flags=re.MULTILINE)  # Remove trailing whitespace
            
            # Create output file path
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            company_name = buyer_info['company_name'].replace('/', '_').replace('\\', '_').replace(':', '_')
            sanitized_name = re.sub(r'[<>:"/\\|?*]', '_', company_name)
            
            # Determine output directory
            output_dir = self.no_contact_dir if not has_contacts else self.output_dir
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Create output file path - Use .eml which is more universally supported
            output_file = output_dir / f"{buyer_type}_Email_{sanitized_name}_{timestamp}.eml"
            
            # Create an .eml file (standard email format)
            with open(output_file, 'w') as f:
                # Add email headers
                f.write(f"From: sender@example.com\n")
                f.write(f"To: {recipients_str}\n")
                f.write(f"Subject: {subject}\n")
                f.write("MIME-Version: 1.0\n")
                f.write("Content-Type: text/plain; charset=utf-8\n")
                f.write("\n")  # Empty line separates headers from body
                f.write(body)
            
            logger.info(f"Successfully created .eml file: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Error creating email file: {str(e)}")
            import traceback
            logger.error(f"Detailed error traceback:\n{traceback.format_exc()}")
            
            # Fall back to docx format
            logger.info("Falling back to docx format")
            return self.save_email_as_docx(email_content, buyer_info, has_contacts, buyer_type)

    def _is_valid_email(self, email: str) -> bool:
        """
        Validate if a string is a properly formatted email address.
        
        Args:
            email: Email address to validate
            
        Returns:
            Boolean indicating if the email is valid
        """
        if not email or '@' not in email:
            return False
        
        # More comprehensive email validation pattern
        pattern = r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email))
    
    def process_file(self, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process a single buyer file.
        
        Args:
            file_info: Dictionary with file information
            
        Returns:
            Updated file information dictionary
        """
        try:
            logger.info(f"Processing file: {file_info['filename']}")
            
            # If path is a string, convert to Path
            file_path = file_info['path']
            if isinstance(file_path, str):
                file_path = Path(file_path)
            
            # Get buyer type from file info or extract from filename
            buyer_type = file_info.get('buyer_type', 'UNKNOWN')
            if buyer_type == 'UNKNOWN' and isinstance(file_path, Path):
                # Extract from filename (first part before underscore)
                buyer_type = file_path.name.split('_')[0]
            
            # Extract buyer information
            buyer_info = self.extract_buyer_info(file_path)
            
            # Check if contact information was found
            has_contacts = True  # Default to true
            if not buyer_info['contacts']:
                # Check if the content has "No contact information found"
                if 'full_content' in buyer_info and "No contact information found" in buyer_info['full_content']:
                    has_contacts = False
                    logger.info(f"No contact information found for {buyer_info['company_name']}")
            
            # Generate email
            email_content = self.generate_email(buyer_info)
            
            # Save as Outlook message with appropriate placement
            output_file = self.save_email_as_outlook_msg(
                email_content, 
                buyer_info,
                has_contacts,
                buyer_type
            )
            
            # Update file information
            file_info.update({
                "processed": True,
                "success": output_file is not None,
                "output_file": str(output_file) if output_file else None,  # Convert Path to string
                "error": None,
                "has_contacts": has_contacts
            })
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error processing file {file_info['filename']}: {str(e)}")
            
            file_info.update({
                "processed": True,
                "success": False,
                "output_file": None,
                "error": str(e)
            })
            
            return file_info
    
    def run(self, single_file=None) -> Dict[str, Any]:
        """
        Run the email generation process for all strong buyer files or a single file.
        
        Args:
            single_file: Optional path to a single file to process
            
        Returns:
            Dictionary with results and statistics
        """
        logger.info("Starting email generation process")
        
        # Check if we're processing just a single file
        if single_file:
            single_file_path = Path(single_file)
            if not single_file_path.exists():
                logger.error(f"Single file not found: {single_file}")
                return {
                    "status": "error",
                    "message": f"File not found: {single_file}",
                    "stats": {"total": 0, "processed": 0, "successful": 0, "failed": 0}
                }
            
            # Create file info for the single file
            company_name = single_file_path.stem.split('_')[1]  # Assuming format STRONG_companyname_...
            file_info = {
                "path": str(single_file_path),  # Convert Path to string
                "company_name": company_name,
                "filename": single_file_path.name,
                "processed": False,
                "success": None,
                "output_file": None,
                "error": None
            }
            
            # Process the single file
            logger.info(f"Processing single file: {single_file_path.name}")
            result = self.process_file(file_info)
            
            # Update statistics
            stats = {"total": 1, "processed": 1, "successful": 1 if result["success"] else 0, "failed": 0 if result["success"] else 1}
            
            logger.info(f"Single file processing complete. Success: {result['success']}")
            
            return {
                "status": "complete",
                "results": [result],
                "stats": stats
            }
        
        # Default behavior - process all strong buyer files
        # List strong buyer files
        files = self.list_strong_buyer_files()
        
        if not files:
            logger.warning("No strong buyer files found to process")
            return {
                "status": "complete",
                "message": "No strong buyer files found to process",
                "stats": {"total": 0, "processed": 0, "successful": 0, "failed": 0}
            }
        
        # Process each file
        results = []
        stats = {"total": len(files), "processed": 0, "successful": 0, "failed": 0}
        
        for file_info in files:
            # Convert Path to string to ensure JSON serialization works
            if isinstance(file_info["path"], Path):
                file_info["path"] = str(file_info["path"])
            
            # Process the file
            result = self.process_file(file_info)
            
            # Ensure output_file is a string for JSON serialization
            if result["output_file"] and isinstance(result["output_file"], Path):
                result["output_file"] = str(result["output_file"])
            
            # Update statistics
            stats["processed"] += 1
            if result["success"]:
                stats["successful"] += 1
            else:
                stats["failed"] += 1
            
            # Add to results
            results.append(result)
        
        # Save summary
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        summary_file = self.output_dir / f"email_generation_summary_{timestamp}.json"
        
        with open(summary_file, "w") as f:
            json.dump({
                "stats": stats,
                "results": results
            }, f, indent=2, default=str)  # Use default=str to handle any other non-serializable objects
        
        logger.info(f"Email generation complete. Stats: {stats}")
        
        return {
            "status": "complete",
            "results": results,
            "stats": stats,
            "summary_file": str(summary_file)  # Convert Path to string
        } 